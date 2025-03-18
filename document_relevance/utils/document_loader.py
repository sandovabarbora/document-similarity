"""
Document loader module for the Document Relevance Classification System.
Handles loading and processing documents from various file formats with support for recursive directory traversal.
"""

import os
import re
import time
import hashlib
import mimetypes
from typing import List, Dict, Tuple, Optional, Any, Union, Set
from pathlib import Path
import logging
import io
import traceback

from document_relevance.models.document import Document
from document_relevance.utils.logging import get_logger, log_execution_time, log_exceptions

# Initialize logger
logger = get_logger("document_loader")

# Try to import document processing libraries
# PDF processing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available. Will use PyPDF2 for PDF processing.")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not available. PDF processing may be limited.")

# DOCX processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX processing will be limited.")

# Excel processing
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl not available. Excel processing will be limited.")

# Try to import tika for fallback document processing
try:
    import tika
    from tika import parser as tika_parser
    tika.initVM()
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.info("Apache Tika not available. Fallback document processing will be limited.")


class DocumentPreprocessor:
    """Class to preprocess document text."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the document preprocessor.
        
        Args:
            config: Configuration options
        """
        self.config = config or {}
        self.clean_pattern = re.compile(r'[\n\r\t]+')
        self.min_line_length = self.config.get('min_line_length', 2)
        self.max_consecutive_spaces = self.config.get('max_consecutive_spaces', 2)
        self.remove_urls = self.config.get('remove_urls', False)
        self.remove_emails = self.config.get('remove_emails', False)
        self.logger = get_logger("document_preprocessor")
        
        # Regular expressions for cleaning
        self.url_pattern = re.compile(r'https?://\S+|www\.\S+')
        self.email_pattern = re.compile(r'\S+@\S+\.\S+')
        
    def clean_text(self, text: str) -> str:
        """
        Clean the text by removing excessive whitespace, etc.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned text
        """
        # Skip if empty
        if not text:
            return ""
            
        # Replace newlines, tabs with spaces
        text = self.clean_pattern.sub(' ', text)
        
        # Remove URLs if configured
        if self.remove_urls:
            text = self.url_pattern.sub('', text)
            
        # Remove emails if configured
        if self.remove_emails:
            text = self.email_pattern.sub('', text)
        
        # Remove multiple spaces (more than max_consecutive_spaces)
        text = re.sub(r'\s{' + str(self.max_consecutive_spaces+1) + ',}', ' ' * self.max_consecutive_spaces, text)
        
        # Remove very short lines that are likely noise
        lines = text.split('\n')
        cleaned_lines = [line for line in lines if len(line.strip()) >= self.min_line_length]
        
        return ' '.join(cleaned_lines).strip()
    
    def preprocess(self, document: Document) -> Document:
        """
        Preprocess the document content.
        
        Args:
            document: Document to preprocess
            
        Returns:
            Preprocessed document
        """
        start_length = len(document.content)
        document.content = self.clean_text(document.content)
        end_length = len(document.content)
        
        self.logger.debug(f"Preprocessed document {document.id}: {start_length} chars -> {end_length} chars")
        
        # Log a warning if preprocessing reduced content too much
        if start_length > 0 and end_length / start_length < 0.5:
            self.logger.warning(
                f"Preprocessing reduced document {document.id} content by more than 50%: "
                f"{start_length} -> {end_length} chars"
            )
            
        return document


class DocumentLoader:
    """
    Class to load documents from files with recursive directory support.
    Handles various file formats and extracts content and metadata.
    """
    
    def __init__(self, preprocessor: Optional[DocumentPreprocessor] = None, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the document loader.
        
        Args:
            preprocessor: Document preprocessor for text cleaning
            config: Configuration options
        """
        self.config = config or {}
        self.preprocessor = preprocessor or DocumentPreprocessor()
        self.logger = get_logger("document_loader")
        self.max_file_size = self.config.get('max_file_size', 100 * 1024 * 1024)  # 100MB default
        
        # Default supported extensions and their loaders
        self.supported_extensions = {
            '.txt': self.load_text_file,
            '.md': self.load_text_file,
            '.pdf': self.load_pdf_file,
            '.docx': self.load_docx_file,
            '.doc': self.load_docx_file,  # May not work well with old .doc format
            '.xlsx': self.load_excel_file,
            '.xls': self.load_excel_file,
            '.csv': self.load_csv_file,
            '.json': self.load_text_file,
            '.xml': self.load_text_file,
            '.html': self.load_text_file,
            '.htm': self.load_text_file
        }
        
        # Add or remove extensions based on configuration
        extra_extensions = self.config.get('extra_extensions', {})
        if extra_extensions:
            for ext, loader_name in extra_extensions.items():
                if hasattr(self, loader_name):
                    self.supported_extensions[ext.lower()] = getattr(self, loader_name)
                    
        disabled_extensions = self.config.get('disabled_extensions', [])
        for ext in disabled_extensions:
            if ext.lower() in self.supported_extensions:
                del self.supported_extensions[ext.lower()]
                
        # Log supported formats
        self.logger.info(f"Supported file formats: {', '.join(self.supported_extensions.keys())}")
        
    @log_execution_time()
    @log_exceptions(reraise=True)
    def load_text_file(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load text from a plain text file.
        
        Args:
            filepath: Path to the text file
            
        Returns:
            Tuple of (content, metadata)
        """
        self.logger.debug(f"Loading text file: {filepath}")
        
        # Detect encoding (could use chardet for better detection)
        encoding = 'utf-8'
        metadata = {
            'encoding': encoding,
            'mime_type': mimetypes.guess_type(filepath)[0] or 'text/plain'
        }
        
        try:
            with open(filepath, 'r', encoding=encoding, errors='replace') as f:
                content = f.read()
                
            return content, metadata
        except UnicodeDecodeError:
            # Try with latin-1 encoding if utf-8 fails
            self.logger.warning(f"UTF-8 decoding failed for {filepath}, trying latin-1")
            with open(filepath, 'r', encoding='latin-1', errors='replace') as f:
                content = f.read()
                
            metadata['encoding'] = 'latin-1'
            return content, metadata
            
    @log_execution_time()
    @log_exceptions(reraise=True)
    def load_pdf_file(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load text from a PDF file using available libraries.
        
        Args:
            filepath: Path to the PDF file
            
        Returns:
            Tuple of (content, metadata)
        """
        self.logger.debug(f"Loading PDF file: {filepath}")
        
        # Check if file size is within limits
        file_size = os.path.getsize(filepath)
        if file_size > self.max_file_size:
            msg = f"PDF file size ({file_size/1024/1024:.2f}MB) exceeds limit ({self.max_file_size/1024/1024:.2f}MB)"
            self.logger.error(msg)
            raise ValueError(msg)
        
        metadata = {
            'mime_type': 'application/pdf',
            'extraction_method': None,
            'page_count': 0
        }
        
        # Try PyMuPDF first (better extraction)
        if PYMUPDF_AVAILABLE:
            try:
                content = ""
                metadata['extraction_method'] = 'pymupdf'
                
                with fitz.open(filepath) as pdf:
                    metadata['page_count'] = pdf.page_count
                    
                    # Extract document metadata
                    if pdf.metadata:
                        for key, value in pdf.metadata.items():
                            if value:
                                metadata[f'pdf_{key.lower()}'] = value
                    
                    # Extract text from each page
                    for i, page in enumerate(pdf):
                        self.logger.debug(f"Extracting page {i+1}/{pdf.page_count}")
                        page_text = page.get_text()
                        content += page_text
                        
                return content, metadata
            except Exception as e:
                self.logger.warning(f"PyMuPDF extraction failed for {filepath}: {e}. Trying alternative method.")
        
        # Try PyPDF2 if PyMuPDF failed or is not available
        if PYPDF2_AVAILABLE:
            try:
                content = ""
                metadata['extraction_method'] = 'pypdf2'
                
                with open(filepath, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(reader.pages)
                    
                    # Extract document metadata
                    if reader.metadata:
                        for key, value in reader.metadata.items():
                            if value:
                                # Remove the leading slash in keys
                                clean_key = key.strip('/').lower()
                                metadata[f'pdf_{clean_key}'] = value
                    
                    # Extract text from each page
                    for i in range(len(reader.pages)):
                        self.logger.debug(f"Extracting page {i+1}/{len(reader.pages)}")
                        page_text = reader.pages[i].extract_text() or ""
                        content += page_text
                        
                return content, metadata
            except Exception as e:
                self.logger.warning(f"PyPDF2 extraction failed for {filepath}: {e}. Trying alternative method.")
        
        # Try Tika as a last resort
        if TIKA_AVAILABLE:
            try:
                parsed = tika_parser.from_file(filepath)
                content = parsed["content"] or ""
                metadata['extraction_method'] = 'tika'
                
                # Add Tika metadata
                if parsed["metadata"]:
                    for key, value in parsed["metadata"].items():
                        if value:
                            metadata[f'tika_{key.lower().replace(" ", "_")}'] = value
                            
                return content, metadata
            except Exception as e:
                self.logger.error(f"All PDF extraction methods failed for {filepath}: {e}")
                
        # If all methods failed and we get here
        self.logger.error(f"No PDF extraction method available for {filepath}")
        return f"[PDF EXTRACTION FAILED: {filepath}]", metadata
    
    @log_execution_time()
    @log_exceptions(reraise=True)
    def load_docx_file(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load text from a DOCX file.
        
        Args:
            filepath: Path to the DOCX file
            
        Returns:
            Tuple of (content, metadata)
        """
        self.logger.debug(f"Loading DOCX file: {filepath}")
        
        metadata = {
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'extraction_method': None
        }
        
        # Try python-docx if available
        if DOCX_AVAILABLE:
            try:
                metadata['extraction_method'] = 'python-docx'
                doc = docx.Document(filepath)
                
                # Extract document properties if available
                try:
                    core_props = doc.core_properties
                    if core_props:
                        metadata['docx_author'] = core_props.author
                        metadata['docx_title'] = core_props.title
                        metadata['docx_subject'] = core_props.subject
                        metadata['docx_created'] = str(core_props.created) if core_props.created else None
                        metadata['docx_modified'] = str(core_props.modified) if core_props.modified else None
                except:
                    self.logger.debug(f"Could not extract core properties from {filepath}")
                
                # Extract text from paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            paragraphs.append(" | ".join(row_text))
                
                content = "\n\n".join(paragraphs)
                metadata['paragraph_count'] = len(doc.paragraphs)
                metadata['table_count'] = len(doc.tables)
                
                return content, metadata
            except Exception as e:
                self.logger.warning(f"python-docx extraction failed for {filepath}: {e}. Trying alternative method.")
        
        # Try Tika as a fallback
        if TIKA_AVAILABLE:
            try:
                parsed = tika_parser.from_file(filepath)
                content = parsed["content"] or ""
                metadata['extraction_method'] = 'tika'
                
                # Add Tika metadata
                if parsed["metadata"]:
                    for key, value in parsed["metadata"].items():
                        if value:
                            metadata[f'tika_{key.lower().replace(" ", "_")}'] = value
                            
                return content, metadata
            except Exception as e:
                self.logger.error(f"All DOCX extraction methods failed for {filepath}: {e}")
        
        # If we get here, no method worked
        self.logger.error(f"No DOCX extraction method available for {filepath}")
        return f"[DOCX EXTRACTION FAILED: {filepath}]", metadata
    
    @log_execution_time()
    @log_exceptions(reraise=True)
    def load_excel_file(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load text from an Excel file.
        
        Args:
            filepath: Path to the Excel file
            
        Returns:
            Tuple of (content, metadata)
        """
        self.logger.debug(f"Loading Excel file: {filepath}")
        
        if not EXCEL_AVAILABLE:
            self.logger.warning("Excel processing not available. Install openpyxl package.")
            return f"[EXCEL PROCESSING NOT AVAILABLE: {filepath}]", {'mime_type': 'application/vnd.ms-excel'}
        
        metadata = {
            'mime_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'extraction_method': 'openpyxl'
        }
        
        try:
            # Load workbook
            workbook = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            metadata['sheet_count'] = len(workbook.sheetnames)
            metadata['sheet_names'] = workbook.sheetnames
            
            # Process each sheet
            all_sheet_content = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = [f"Sheet: {sheet_name}"]
                
                # Process rows
                row_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None and str(cell).strip() for cell in row):
                        row_data.append(" | ".join(str(cell) if cell is not None else "" for cell in row))
                
                if row_data:
                    sheet_content.extend(row_data)
                    all_sheet_content.append("\n".join(sheet_content))
            
            content = "\n\n".join(all_sheet_content)
            return content, metadata
        except Exception as e:
            self.logger.error(f"Excel extraction failed for {filepath}: {e}")
            return f"[EXCEL EXTRACTION FAILED: {filepath}]", metadata
    
    @log_execution_time()
    @log_exceptions(reraise=True)
    def load_csv_file(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load text from a CSV file.
        
        Args:
            filepath: Path to the CSV file
            
        Returns:
            Tuple of (content, metadata)
        """
        self.logger.debug(f"Loading CSV file: {filepath}")
        import csv
        
        metadata = {
            'mime_type': 'text/csv',
            'extraction_method': 'csv'
        }
        
        try:
            # Try to detect encoding
            encoding = 'utf-8'
            
            with open(filepath, 'r', encoding=encoding, errors='replace', newline='') as f:
                # Try to detect delimiter
                sample = f.read(4096)
                f.seek(0)
                
                dialect = csv.Sniffer().sniff(sample)
                metadata['delimiter'] = dialect.delimiter
                
                # Read CSV
                reader = csv.reader(f, dialect)
                rows = [" | ".join(cell for cell in row) for row in reader]
                
                metadata['row_count'] = len(rows)
                content = "\n".join(rows)
                
                return content, metadata
        except Exception as e:
            self.logger.warning(f"CSV parsing failed for {filepath}: {e}. Falling back to text reading.")
            
            # Fall back to simple text reading
            try:
                with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                return content, metadata
            except Exception as e2:
                self.logger.error(f"Text fallback also failed for {filepath}: {e2}")
                return f"[CSV EXTRACTION FAILED: {filepath}]", metadata
    
    def get_document_title(self, filepath: str, content: str = "", metadata: Dict[str, Any] = None) -> str:
        """
        Try to extract a meaningful title from the document.
        
        Args:
            filepath: Path to the document
            content: Document content
            metadata: Document metadata
            
        Returns:
            Document title
        """
        # First check if we have a title in metadata
        if metadata:
            for key in ['pdf_title', 'docx_title', 'tika_title']:
                if key in metadata and metadata[key]:
                    return metadata[key]
        
        # Otherwise use the filename without extension
        filename = os.path.basename(filepath)
        title = os.path.splitext(filename)[0]
        
        # Try to improve title if it's just a number or very short
        if title.isdigit() or len(title) < 3:
            # Use parent folder name + filename
            parent_folder = os.path.basename(os.path.dirname(filepath))
            if parent_folder:
                title = f"{parent_folder} - {title}"
        
        # If we have content, try to use the first non-empty line
        if content and len(content) > 10:
            first_lines = content.strip().split('\n')
            for line in first_lines:
                line = line.strip()
                if line and len(line) > 3 and len(line) < 100:  # Reasonable title length
                    return line
                    
        return title
    
    def calculate_document_id(self, filepath: str, reference_dir: Optional[str] = None) -> str:
        """
        Calculate a document ID based on its filepath.
        
        Args:
            filepath: Path to the document
            reference_dir: Reference directory for relative path calculation
            
        Returns:
            Document ID
        """
        if reference_dir and os.path.exists(reference_dir):
            # Use relative path from reference directory
            try:
                rel_path = os.path.relpath(filepath, reference_dir)
                return rel_path
            except:
                pass
                
        # Fallback: use filename
        filename = os.path.basename(filepath)
        
        # Add parent folder for better identification
        parent = os.path.basename(os.path.dirname(filepath))
        if parent:
            return f"{parent}/{filename}"
            
        return filename
    
    @log_execution_time()
    @log_exceptions(reraise=True)
    def load_document(self, filepath: str, is_reference: bool = False) -> Document:
        """
        Load a document from a file based on its extension.
        
        Args:
            filepath: Path to the document
            is_reference: Whether this is a reference document
            
        Returns:
            Loaded and preprocessed Document
        """
        self.logger.info(f"Loading document: {filepath}")
        start_time = time.time()
        
        # Check if file exists
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
            
        # Get reference directory from config
        reference_dir = self.config.get('reference_dir')
        
        # Calculate document ID (relative path or filename)
        doc_id = self.calculate_document_id(filepath, reference_dir)
        
        # Determine file extension
        _, ext = os.path.splitext(filepath.lower())
        
        # Check if extension is supported
        if ext not in self.supported_extensions:
            supported = ", ".join(self.supported_extensions.keys())
            raise ValueError(f"Unsupported file format: {ext}. Supported formats: {supported}")
        
        # Load content based on file extension
        loader_func = self.supported_extensions[ext]
        try:
            content, file_metadata = loader_func(filepath)
        except Exception as e:
            self.logger.error(f"Error loading document content: {e}")
            traceback.print_exc()
            raise
        
        # Get document title
        title = self.get_document_title(filepath, content, file_metadata)
        
        # Calculate content hash for identifying duplicates
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
        
        # Add file system metadata
        metadata = {
            'file_size': os.path.getsize(filepath),
            'file_type': ext[1:],  # Extension without dot
            'last_modified': os.path.getmtime(filepath),
            'creation_time': os.path.getctime(filepath) if hasattr(os, 'getctime') else None,
            'char_count': len(content),
            'relative_path': os.path.relpath(filepath, reference_dir) if reference_dir else None,
            'parent_folder': os.path.basename(os.path.dirname(filepath)),
            'content_hash': content_hash,
            'filename': os.path.basename(filepath)
        }
        
        # Add file-specific metadata
        metadata.update(file_metadata)
        
        # Create document
        doc = Document(
            id=doc_id,
            title=title,
            content=content,
            filepath=filepath,
            metadata=metadata,
            is_reference=is_reference
        )
        
        # Preprocess content
        doc = self.preprocessor.preprocess(doc)
        
        elapsed = time.time() - start_time
        self.logger.info(f"Loaded document {doc_id} ({len(doc.content)} chars) in {elapsed:.2f}s")
        
        return doc
    
    @log_execution_time()
    def load_documents_from_directory(
        self, 
        directory: str, 
        is_reference: bool = False,
        recursive: bool = True,
        extensions: Optional[List[str]] = None
    ) -> List[Document]:
        """
        Load all documents from a directory and its subdirectories.
        
        Args:
            directory: Path to the directory
            is_reference: Whether these are reference documents
            recursive: Whether to process subdirectories recursively
            extensions: List of file extensions to process (if None, use all supported)
            
        Returns:
            List of loaded documents
        """
        documents = []
        self.logger.info(f"Loading documents from {directory} (recursive={recursive})")
        
        # Check if directory exists
        if not os.path.exists(directory):
            self.logger.error(f"Directory does not exist: {directory}")
            return documents
        
        # Store the reference directory for relative path calculation
        self.config['reference_dir'] = directory
        
        # Filter extensions if provided
        supported_exts = self.supported_extensions.keys()
        if extensions:
            # Ensure extensions start with a dot
            extensions = [ext if ext.startswith('.') else f'.{ext}' for ext in extensions]
            # Only keep extensions that are both supported and requested
            exts_to_process = [ext for ext in extensions if ext.lower() in supported_exts]
            self.logger.info(f"Filtering to extensions: {', '.join(exts_to_process)}")
        else:
            exts_to_process = supported_exts
        
        # Get all files to process
        files_to_process = []
        
        if recursive:
            # Walk the directory tree recursively
            for root, _, files in os.walk(directory):
                for filename in files:
                    filepath = os.path.join(root, filename)
                    _, ext = os.path.splitext(filename.lower())
                    if ext in exts_to_process:
                        files_to_process.append(filepath)
        else:
            # Only process files in the top directory
            for filename in os.listdir(directory):
                filepath = os.path.join(directory, filename)
                if os.path.isfile(filepath):
                    _, ext = os.path.splitext(filename.lower())
                    if ext in exts_to_process:
                        files_to_process.append(filepath)
        
        # Log summary
        self.logger.info(f"Found {len(files_to_process)} documents to process in {directory}")
        
        # Process each file
        for i, filepath in enumerate(files_to_process):
            try:
                self.logger.info(f"Processing document {i+1}/{len(files_to_process)}: {filepath}")
                doc = self.load_document(filepath, is_reference)
                documents.append(doc)
                self.logger.info(f"Successfully loaded document: {filepath}")
            except Exception as e:
                self.logger.warning(f"Failed to load document {filepath}: {e}")
        
        self.logger.info(f"Successfully loaded {len(documents)} out of {len(files_to_process)} documents")
        
        return documents
    
    def detect_duplicates(self, documents: List[Document]) -> List[List[Document]]:
        """
        Detect duplicate documents based on content hash.
        
        Args:
            documents: List of documents to check
            
        Returns:
            List of lists containing duplicate documents
        """
        # Group documents by content hash
        hash_groups = {}
        for doc in documents:
            content_hash = doc.metadata.get('content_hash')
            if content_hash:
                if content_hash in hash_groups:
                    hash_groups[content_hash].append(doc)
                else:
                    hash_groups[content_hash] = [doc]
        
        # Return groups with more than one document (duplicates)
        duplicate_groups = [group for hash_val, group in hash_groups.items() if len(group) > 1]
        
        if duplicate_groups:
            self.logger.info(f"Found {len(duplicate_groups)} groups of duplicate documents")
            for i, group in enumerate(duplicate_groups):
                self.logger.debug(f"Duplicate group {i+1}: {', '.join(doc.id for doc in group)}")
        
        return duplicate_groups